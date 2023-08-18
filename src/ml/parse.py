import os
import docx
from ppocr import PPOcr
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LAParams, LTTextBoxHorizontal
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfinterp import PDFPageInterpreter, PDFResourceManager
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfparser import PDFParser
from win32com import client
from ppocr.model import Success, Error


def parse_image(ocr: PPOcr, file_path: str):
    result = ocr.run(file_path)
    ans=""
    # match result:
    #     case Success():
            
    for line in result:
        for i in line:
            temp = i[1][0]
            ans = ans + temp
    with open(f"{file_path}.txt", "w", encoding='utf-8') as f:
        f.write(ans)


def parse_pdf_file(path, filename, txt_save_path):
    """
    功能：解析pdf 文本，保存到txt文件中
    path：pdf存放的文件夹路径
    filename: pdf文件名
    txt_save_path: 需要保存的txt文件夹路径
    """
    pdf_path=os.path.join(path,filename)
    fp = open(pdf_path, 'rb') # 以二进制读模式打开
    #用文件对象来创建一个pdf文档分析器
    praser = PDFParser(fp)
    # 创建一个PDF文档
    doc = PDFDocument(praser)
    # 连接分析器 与文档对象
    praser.set_document(doc)

    # 检测文档是否提供txt转换，不提供就忽略
    if doc.is_extractable:
        # 创建PDf 资源管理器 来管理共享资源
        rsrcmgr = PDFResourceManager()
        # 创建一个PDF设备对象
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        # 循环遍历列表，每次处理一个page的内容
        for page in PDFPage.create_pages(doc): # doc.get_pages() 获取page列表
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等 想要获取文本就获得对象的text属性，
            for x in layout:
                if (isinstance(x, LTTextBoxHorizontal)):
                    with open(os.path.join(txt_save_path, f'{filename}.txt'), 'a', encoding='utf-8') as f:
                        results = x.get_text()
                        print(results)
                        results=results.replace("\n","")
                        f.write(results + '\n')


def parse_docx_file(file_path: str):
    docx_reader = docx.Document(file_path)
    with open(f"{file_path}.txt", "w", encoding='utf-8') as f:
        for par in docx_reader.paragraphs:
            f.write(par.text)
            f.write("\n")


def parse_doc_file(file_path: str):
    path = file_path.replace("/", "\\")
    word = client.DispatchEx("Word.Application")
    print(path)
    doc = word.Documents.Open(path)
    tempname = f"{path}x"
    doc.SaveAs(tempname, 12)

    doc.Close()
    parse_docx_file(file_path=tempname)