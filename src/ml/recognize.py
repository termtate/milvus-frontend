from pprint import pprint
from injector import inject
import os
from logging import getLogger
import docx
from win32com import client
from pdfminer.pdfparser import PDFParser
from pdfminer.pdfdocument import PDFDocument
from pdfminer.pdfpage import PDFPage
from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
from pdfminer.converter import PDFPageAggregator
from pdfminer.layout import LTTextBoxHorizontal,LAParams
from paddleocr import PaddleOCR
import torch
from ml.utils import get_label, get_vocab, extract, label_list, read_info
from ml.config import settings
from ml.model import Model

logger = getLogger(__name__)


class Recognizer:
    @inject
    def __init__(self, paddleocr: PaddleOCR):
      self.paddleocr = paddleocr
    
    def read_files2(self, path):
        filename_list = os.listdir(path)
        for filename in filename_list:
            self._recognize_from_file(path, filename)

        files=os.listdir(path)
        for file in files:
            post_name=file.split(".")[-1]
            if post_name != "txt":
                continue
            self.read_generated_txt(os.path.join(path, file))
    
    
    def read_generated_txt(self, path: str):
        with open(path, "r", encoding="UTF-8") as f:
            text = f.read()
            self.output_labels(text)
                
        # 输出标签
    def output_labels(self, text):
        _, word2id = get_vocab()
        input_ = torch.tensor([[word2id.get(w, settings.WORD_UNK_ID) for w in text]]).to(settings.DEVICE)
        mask = torch.tensor([[1] * len(text)]).bool().to(settings.DEVICE)

        model = torch.load(f'{settings.MODEL_DIR}model_7000.pth', map_location=settings.DEVICE)
        y_pred = model(input_, mask)
        id2label, _ = get_label()

        label = [id2label[l] for l in y_pred[0]]
        # print(text)
        # print(label)
        info = extract(label, text)

        # pprint("------info-------")
        # pprint(info)
        # pprint('\n' * 3)
        # pprint("------text-------")
        # pprint(info)
        # print(info)
        # load_db(info, text)  # TODO
        # print()
        labels = label_list()
        # print(labels)
        output = read_info(text)
        # pprint(labels)
        for l in labels:
            output.append([l, ''])
        for i, label_name in enumerate(labels, start=20):
            for per in info:
                if isinstance(per, list) and per[0] == label_name:
                    if output[i][1] != '':
                        output[i][1] += '，'
                    output[i][1] += per[1]
        pprint(dict(output))

        sqlstr = list(map(str, labels))
        sqlstr = ','.join(sqlstr)
        sqlstr = f'第几次住院,姓名,病案号,性别,年龄,电话,发作演变过程,发作持续时间,发作频次,母孕年龄,孕次产出,出生体重,头围,血、尿代谢筛查,铜兰蛋白,脑脊液,基因检查,头部CT,头部MRI,头皮脑电图,{sqlstr}'
        # print(sqlstr)

        aaa = ''
        for i in range(78):
            if i != 0:
                aaa += ','
            aaa += '%s'

        sql = f'INSERT INTO TABLE1({sqlstr}) VALUE ({aaa})'
        value = [v[1] for v in output if isinstance(v, list)]
        value = tuple(value)
        # pprint(sql)
        # pprint(value)


    def _recognize_from_file(self, pathname: str, filename: str):
        '''为非txt格式的文件，在相同路径下生成对应的txt文件，命名为原文件名+.txt'''
        complete_name = os.path.join(pathname, filename)

        file_type = filename.split(".")[-1]

        match file_type:
            case "txt":
                return
            case "docx":
                _parse_docx_file(complete_name)
            
            case "doc":
                _parse_doc_file(complete_name)
                
            case "pdf":
                _parse_pdf_file(pathname, filename, pathname)
                
            case "png" | "jpg":
                self._parse_image(complete_name)
            
            case _:
                logger.warning("Not Support File Type")
    
    
    def _parse_image(self, file_path: str):
        result = self.paddleocr.ocr(file_path, cls=True)
        ans=""
        for line in result:
            for i in line:
                temp = i[1][0]
                ans = ans + temp
        with open(f"{file_path}.txt", "w", encoding='utf-8') as f:
            f.write(ans)


def _parse_pdf_file(path, filename, txt_save_path):
    """
    功能：解析pdf 文本，保存到txt文件中
    path：pdf存放的文件夹路径
    filename: pdf文件名
    txt_save_path: 需要保存的txt文件夹路径

    """
    pdf_path=os.path.join(path,filename)
    fp = open(pdf_path, 'rb') # 以二进制读模式打开
    #用文件对象来创建一个pdf文档分析器
    praser = PDFParser(fp)
    # 创建一个PDF文档
    doc = PDFDocument(praser)
    # 连接分析器 与文档对象
    praser.set_document(doc)

    # 检测文档是否提供txt转换，不提供就忽略
    if doc.is_extractable:
        # 创建PDf 资源管理器 来管理共享资源
        rsrcmgr = PDFResourceManager()
        # 创建一个PDF设备对象
        laparams = LAParams()
        device = PDFPageAggregator(rsrcmgr, laparams=laparams)
        # 创建一个PDF解释器对象
        interpreter = PDFPageInterpreter(rsrcmgr, device)

        # 循环遍历列表，每次处理一个page的内容
        for page in PDFPage.create_pages(doc): # doc.get_pages() 获取page列表
            interpreter.process_page(page)
            # 接受该页面的LTPage对象
            layout = device.get_result()
            # 这里layout是一个LTPage对象 里面存放着 这个page解析出的各种对象 一般包括LTTextBox, LTFigure, LTImage, LTTextBoxHorizontal 等等 想要获取文本就获得对象的text属性，
            for x in layout:
                if (isinstance(x, LTTextBoxHorizontal)):
                    with open(os.path.join(txt_save_path, f'{filename}.txt'), 'a', encoding='utf-8') as f:
                        results = x.get_text()
                        print(results)
                        results=results.replace("\n","")
                        f.write(results + '\n')

def _parse_docx_file(file_path: str):
    docx_reader = docx.Document(file_path)
    with open(f"{file_path}.txt", "w", encoding='utf-8') as f:
        for par in docx_reader.paragraphs:
            f.write(par.text)
            f.write("\n")


def _parse_doc_file(file_path: str):
    word = client.DispatchEx("Word.Application")
    print(file_path)
    doc = word.Documents.Open(file_path)
    tempname = f"{file_path}x"
    doc.SaveAs(tempname, 12)
    
    doc.Close()
    _parse_docx_file(file_path=tempname)